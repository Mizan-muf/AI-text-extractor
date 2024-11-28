import os
import time
import re
import google.generativeai as genai
from docx import Document  # Import the Document class from python-docx
import tkinter as tk
from tkinter import filedialog

# Configure Gemini API
genai.configure(api_key="Your_API_here")

def upload_to_gemini(path, mime_type=None):
    """Uploads the given file to Gemini."""
    file = genai.upload_file(path, mime_type=mime_type)
    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
    return file

def wait_for_files_active(files):
    """Waits for the given files to be active."""
    print("Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai.get_file(name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(10)
            file = genai.get_file(name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
    print("...all files ready\n")

# Function to open folder dialog and get the file path
def select_input_file():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename(title="Select a PDF file", filetypes=[("PDF Files", "*.pdf")])
    return file_path

# Function to open folder dialog and get the output directory path
def select_output_folder():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    folder_path = filedialog.askdirectory(title="Select Output Folder")
    return folder_path

# Create the model
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
)

# Let the user select the input PDF file
pdf_file_path = select_input_file()
if not pdf_file_path:
    print("No file selected. Exiting.")
    exit()

# Let the user select the output folder
output_folder = select_output_folder()
if not output_folder:
    print("No output folder selected. Exiting.")
    exit()

# Upload the PDF file to Gemini
files = [upload_to_gemini(pdf_file_path, mime_type="application/pdf")]

# Wait for the uploaded file to be ready
wait_for_files_active(files)

# Prepare the chat session with the specified prompt
chat_session = model.start_chat(
    history=[{
        "role": "user",
        "parts": [
            files[0],
            """You are an AI assistant tasked with processing text. so your role is to correct any spelling, grammar, or formatting issues in the text.
            Correct any spelling or grammar mistakes you encounter.
            If the text contains mathematical expressions, equations, or code snippets, format them using appropriate docx code blocks with proper syntax highlighting.
            Organize the content with appropriate headings and subheadings using docx syntax.
            Ensure that the output is well-formatted and easy to read.""",
        ],
    }]
)

# Send the message and get the response
response = chat_session.send_message("Analyze and format the PDF content.")

# Construct the output file path
output_file_path = os.path.join(output_folder, "result.docx")

# Parse and save the response to a Word file
doc = Document()  # Create a new Document
doc.add_heading('Processed Text from PDF', level=1)  # Add a main title

# Parsing the response text
lines = response.text.splitlines()
for line in lines:
    if line.startswith("### "):  # Subheading level
        doc.add_heading(line[3:], level=2)
    elif line.startswith("## "):  # Heading level
        doc.add_heading(line[2:], level=2)
    elif line.startswith("# "):  # Heading level
        doc.add_heading(line[2:], level=1)
    else:
        # Handle bold text within paragraphs
        paragraph = doc.add_paragraph()
        bold_parts = re.split(r"(\*\*[^*]+\*\*)", line)
        for part in bold_parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True  # Bold text
            else:
                paragraph.add_run(part)  # Regular text

# Save the document
doc.save(output_file_path)
print(f"Response saved to {output_file_path}")
import os
import time
import re
import google.generativeai as genai
from docx import Document  # Import the Document class from python-docx
import tkinter as tk
from tkinter import filedialog

# Configure Gemini API
genai.configure(api_key="Your_API_here")

def upload_to_gemini(path, mime_type=None):
    """Uploads the given file to Gemini."""
    file = genai.upload_file(path, mime_type=mime_type)
    print(f"Uploaded file '{file.display_name}' as: {file.uri}")
    return file

def wait_for_files_active(files):
    """Waits for the given files to be active."""
    print("Waiting for file processing...")
    for name in (file.name for file in files):
        file = genai.get_file(name)
        while file.state.name == "PROCESSING":
            print(".", end="", flush=True)
            time.sleep(10)
            file = genai.get_file(name)
        if file.state.name != "ACTIVE":
            raise Exception(f"File {file.name} failed to process")
    print("...all files ready\n")

# Function to open folder dialog and get the file path
def select_input_file():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    file_path = filedialog.askopenfilename(title="Select a PDF file", filetypes=[("PDF Files", "*.pdf")])
    return file_path

# Function to open folder dialog and get the output directory path
def select_output_folder():
    root = tk.Tk()
    root.withdraw()  # Hide the root window
    folder_path = filedialog.askdirectory(title="Select Output Folder")
    return folder_path

# Create the model
generation_config = {
    "temperature": 1,
    "top_p": 0.95,
    "top_k": 40,
    "max_output_tokens": 8192,
    "response_mime_type": "text/plain",
}

model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=generation_config,
)

# Let the user select the input PDF file
pdf_file_path = select_input_file()
if not pdf_file_path:
    print("No file selected. Exiting.")
    exit()

# Let the user select the output folder
output_folder = select_output_folder()
if not output_folder:
    print("No output folder selected. Exiting.")
    exit()

# Upload the PDF file to Gemini
files = [upload_to_gemini(pdf_file_path, mime_type="application/pdf")]

# Wait for the uploaded file to be ready
wait_for_files_active(files)

# Prepare the chat session with the specified prompt
chat_session = model.start_chat(
    history=[{
        "role": "user",
        "parts": [
            files[0],
            """You are an AI assistant tasked with processing text. so your role is to correct any spelling, grammar, or formatting issues in the text.
            Correct any spelling or grammar mistakes you encounter.
            If the text contains mathematical expressions, equations, or code snippets, format them using appropriate docx code blocks with proper syntax highlighting.
            Organize the content with appropriate headings and subheadings using docx syntax.
            Ensure that the output is well-formatted and easy to read.""",
        ],
    }]
)

# Send the message and get the response
response = chat_session.send_message("Analyze and format the PDF content.")

# Construct the output file path
output_file_path = os.path.join(output_folder, "result.docx")

# Parse and save the response to a Word file
doc = Document()  # Create a new Document
doc.add_heading('Processed Text from PDF', level=1)  # Add a main title

# Parsing the response text
lines = response.text.splitlines()
for line in lines:
    if line.startswith("### "):  # Subheading level
        doc.add_heading(line[3:], level=2)
    elif line.startswith("## "):  # Heading level
        doc.add_heading(line[2:], level=2)
    elif line.startswith("# "):  # Heading level
        doc.add_heading(line[2:], level=1)
    else:
        # Handle bold text within paragraphs
        paragraph = doc.add_paragraph()
        bold_parts = re.split(r"(\*\*[^*]+\*\*)", line)
        for part in bold_parts:
            if part.startswith("**") and part.endswith("**"):
                paragraph.add_run(part[2:-2]).bold = True  # Bold text
            else:
                paragraph.add_run(part)  # Regular text

# Save the document
doc.save(output_file_path)
print(f"Response saved to {output_file_path}")
